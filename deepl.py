import time
import clipboard
from selenium import webdriver
from selenium.webdriver.common.by import By
import os
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.proxy import Proxy, ProxyType
from selenium.webdriver.chrome.options import Options
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from selenium.webdriver.common.desired_capabilities import DesiredCapabilities
from docx import Document

caps = DesiredCapabilities().CHROME
caps["pageLoadStrategy"] = "none"

def detect_docx():
	while True:
		for file in os.listdir("Translated"):
			if file.endswith(".docx"):
				return file

def on_created(event):
	download_file_path = 'Translated/'+detect_docx()
	print(download_file_path)
	f = open(download_file_path, 'rb')
	document = Document(f)
	f.close()
	doc = Document()
	for par in document.paragraphs:
		if len(par.text) > 0:
			doc.add_paragraph(par.text)
	os.remove(download_file_path)
	doc.save(download_file_path)
	observer.stop()
	driver.get("https://papersowl.com/free-plagiarism-checker")
	WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.CSS_SELECTOR, "#terms_box")))
	driver.execute_script("document.querySelector('#terms_box').checked = true;")
	upload_file = WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.CSS_SELECTOR, ".dz-hidden-input")))
	upload_file.send_keys(os.path.abspath("Translated/English uk.docx"))
	WebDriverWait(driver, 6000).until(EC.presence_of_element_located((By.CSS_SELECTOR, ".js_checker_module__result.show")))
	originals_percentage_string = driver.find_element(By.CSS_SELECTOR, '.js_checker_module__original-percent').text
	originals_percentage = float(originals_percentage_string.strip('%'));
	if originals_percentage < 80:
		print("fail", originals_percentage)
	else: print("complete")


options = Options() 
options.add_extension('zenmate.crx')
# options.add_argument("--headless")
options.add_experimental_option("prefs", {
  "download.default_directory": os.path.abspath("Translated")
  })

driver = webdriver.Chrome(options=options, desired_capabilities=caps)


#---------------------VPN--------------------------------------------------
driver.get('chrome-extension://fdcgdnkidjaadafnichfpabhfomcebme/index.html')
while True:
	print(driver.current_url)
	driver.switch_to.window(driver.window_handles[-1])
	if driver.current_url.find('chrome-extension') == -1:
		driver.close()
		driver.switch_to.window(driver.window_handles[0])
		break
file_css = '.img-section img'
WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.CSS_SELECTOR, file_css)))
 # = driver.find_element(By.CSS_SELECTOR, file_css)
driver.execute_script("document.querySelectorAll('.img-section img')[0].click()")
driver.execute_script("document.querySelectorAll('.img-section img')[1].click()")
driver.execute_script("document.querySelectorAll('.img-section img')[2].click()")
driver.execute_script("document.querySelectorAll('.img-section img')[3].click()")


WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.CSS_SELECTOR, '.location-container'))).click()
WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.CSS_SELECTOR, '#country-browsing-US'))).click()
#---------------------VPN--------------------------------------------------


# vpn = EC.find_element((By.CSS_SELECTOR, '.shield-container'))
# time.sleep(2)
# vpn.click()
# Reach the deepL website
deepl_url = 'https://www.deepl.com/fr/translator/files'
driver.get(deepl_url)

#Upload File
file_css = 'input#file-upload_input'
input_file = WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.CSS_SELECTOR, file_css)))
input_file.send_keys(os.path.abspath("English.docx"))

WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.CSS_SELECTOR, 'button[data-testid=doctrans-target-lang-btn]')))
driver.execute_script("document.querySelector('button[data-testid=doctrans-target-lang-btn]').click()")
WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.CSS_SELECTOR, 'button[dl-test=document-translator-lang-option-uk')))
driver.execute_script("document.querySelector('button[dl-test=document-translator-lang-option-uk').click()")

#Submit transalete
input_css = 'button[dl-test=doctrans-translation-button]'
WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.CSS_SELECTOR, input_css)))
driver.execute_script("document.querySelector('button[dl-test=doctrans-translation-button]').click()")

event_handler = FileSystemEventHandler()
event_handler.on_created = on_created


path = os.path.abspath("Translated")
observer = Observer()
observer.schedule(event_handler, path, recursive=True)
observer.start()
try:
	print("Monitoring")
	while True:
		time.sleep(1)
finally:
	print("Stop")
	observer.stop()
	observer.join()
driver.quit()